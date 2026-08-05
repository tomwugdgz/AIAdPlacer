"""青柠智能助手 — 文档生成技能（真实链路）。

文档生成属于「真实链路」：基于真实查询结果落地为可下载的 .docx / .xlsx 文件，
**不是**模拟。依赖 ``python-docx`` / ``openpyxl``（已在 requirements.txt 声明），
采用**懒加载**，模块导入阶段不触碰这些可选依赖，保证服务即使未安装也能正常启动。

技能通过 ``SkillRegistry`` 统一注册与调用。
"""

from __future__ import annotations

import os
from typing import Any, Dict, List, Optional

from app.config import settings
from app.common import setup_logging

logger = setup_logging("qinglin_skills")

# 生成的文档统一落地目录
GENERATED_DIR = os.path.join(os.path.dirname(settings.QINGLIN_MEMORY_DB_PATH), "generated")
os.makedirs(GENERATED_DIR, exist_ok=True)


class DocxSkill:
    """生成 Word 文档（.docx）。"""

    name = "docx"
    description = "生成 Word 文档"

    def run(self, context: Dict[str, Any]) -> Dict[str, Any]:
        from docx import Document
        from docx.shared import Pt

        title = context.get("title", "青柠智能助手文档")
        sections = context.get("sections", []) or []
        table = context.get("table")

        doc = Document()
        doc.add_heading(title, level=0)

        for sec in sections:
            if isinstance(sec, dict):
                heading = sec.get("heading")
                body = sec.get("body", "")
                if heading:
                    doc.add_heading(heading, level=1)
                if body:
                    doc.add_paragraph(body)
            elif isinstance(sec, str):
                doc.add_paragraph(sec)

        if isinstance(table, dict) and table.get("headers") and table.get("rows"):
            headers = table["headers"]
            rows = table["rows"]
            t = doc.add_table(rows=1, cols=len(headers))
            t.style = "Light Grid Accent 1"
            hdr = t.rows[0].cells
            for i, h in enumerate(headers):
                hdr[i].text = str(h)
            for row in rows:
                cells = t.add_row().cells
                for i, val in enumerate(row):
                    cells[i].text = str(val)

        filename = context.get("filename") or f"{title}_{_ts()}.docx"
        filepath = os.path.join(GENERATED_DIR, _safe_name(filename))
        doc.save(filepath)
        logger.info("已生成 DOCX: %s", filepath)
        return {"success": True, "file_path": filepath, "file_name": os.path.basename(filepath), "skill": "docx"}


class XlsxSkill:
    """生成 Excel 文档（.xlsx）。"""

    name = "xlsx"
    description = "生成 Excel 表格"

    def run(self, context: Dict[str, Any]) -> Dict[str, Any]:
        from openpyxl import Workbook

        title = context.get("title", "青柠智能助手表格")
        sheet_name = context.get("sheet_name", "Sheet1")
        headers = context.get("headers", []) or []
        rows = context.get("rows", []) or []

        wb = Workbook()
        ws = wb.active
        ws.title = sheet_name[:31]
        if headers:
            ws.append([str(h) for h in headers])
        for row in rows:
            ws.append([_cell_str(c) for c in row])

        filename = context.get("filename") or f"{title}_{_ts()}.xlsx"
        filepath = os.path.join(GENERATED_DIR, _safe_name(filename))
        wb.save(filepath)
        logger.info("已生成 XLSX: %s", filepath)
        return {"success": True, "file_path": filepath, "file_name": os.path.basename(filepath), "skill": "xlsx"}


class SkillRegistry:
    """文档技能注册表。"""

    def __init__(self):
        self._skills: Dict[str, Any] = {}
        self.register(DocxSkill())
        self.register(XlsxSkill())

    def register(self, skill: Any) -> None:
        self._skills[skill.name] = skill

    def list_skills(self) -> List[str]:
        return list(self._skills.keys())

    def run(self, skill_name: str, context: Dict[str, Any]) -> Dict[str, Any]:
        skill = self._skills.get(skill_name)
        if not skill:
            return {"success": False, "error": f"未知技能：{skill_name}", "available": self.list_skills()}
        try:
            return skill.run(context)
        except ImportError as e:
            msg = (
                f"文档生成依赖缺失（{skill_name}）：{e}。"
                "请执行 `pip install python-docx openpyxl` 后重试。"
            )
            logger.error(msg)
            return {"success": False, "error": msg}
        except Exception as e:  # noqa: BLE001
            logger.exception("文档生成失败")
            return {"success": False, "error": f"文档生成失败：{e}"}


def _ts() -> str:
    from datetime import datetime

    return datetime.now().strftime("%Y%m%d_%H%M%S")


def _safe_name(name: str) -> str:
    import re

    name = re.sub(r'[\\/:*?"<>|]', "_", name)
    return name


def _cell_str(value: Any) -> str:
    if value is None:
        return ""
    return str(value)


# 模块级单例
skill_registry = SkillRegistry()
